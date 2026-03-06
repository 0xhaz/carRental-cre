#!/usr/bin/env python3
"""
Generate RegShield Demo Walkthrough Script (.docx)
Run: python3 generate-demo-script.py
Output: RegShield_Demo_Script.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

# ── Helper Functions ────────────────────────────────────────────────────────

def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

def add_section(text):
    doc.add_heading(text, level=1)

def add_subsection(text):
    doc.add_heading(text, level=2)

def add_step(number, title, time_stamp):
    p = doc.add_paragraph()
    run = p.add_run(f"STEP {number}: {title}")
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run(f"  [{time_stamp}]")
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.size = Pt(10)

def add_action(text):
    """What the presenter DOES (green arrow)"""
    p = doc.add_paragraph()
    run = p.add_run("  DO  ")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.highlight_color = 4  # bright green
    run = p.add_run(f"  {text}")
    run.font.size = Pt(11)

def add_say(text):
    """What the presenter SAYS (blue speech)"""
    p = doc.add_paragraph()
    run = p.add_run("  SAY  ")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.highlight_color = 9  # blue
    run = p.add_run(f'  "{text}"')
    run.italic = True
    run.font.size = Pt(11)

def add_show(text):
    """What appears on screen (yellow eye)"""
    p = doc.add_paragraph()
    run = p.add_run("  SHOW  ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.highlight_color = 7  # yellow
    run = p.add_run(f"  {text}")
    run.font.size = Pt(11)

def add_tip(text):
    """Production tip (gray)"""
    p = doc.add_paragraph()
    run = p.add_run(f"  TIP: {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)

def add_terminal(text):
    """Terminal command"""
    p = doc.add_paragraph()
    run = p.add_run("  TERMINAL  ")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.highlight_color = 1  # black
    run = p.add_run(f"  {text}")
    run.font.name = "Courier New"
    run.font.size = Pt(10)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f"  NOTE: {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 80, 0)
    run.bold = True

def add_blank():
    doc.add_paragraph("")

def add_divider():
    p = doc.add_paragraph()
    run = p.add_run("─" * 70)
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(8)


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT START
# ═══════════════════════════════════════════════════════════════════════════

add_title("RegShield Demo Walkthrough Script")
add_subtitle("Tokenized Vehicle Rental Platform")
add_subtitle("Powered by Chainlink CRE + WorldID + ERC-3643")
add_subtitle(f"Version 1.1 — {datetime.date.today().strftime('%B %d, %Y')}")

add_blank()

# ── Pre-Demo Checklist ──────────────────────────────────────────────────────

add_section("Pre-Demo Checklist")

checklist = [
    "Frontend running locally (npm run dev) at http://localhost:3000",
    "Backend server running (node server.js)",
    "MetaMask installed with Sepolia testnet selected",
    "3 browser profiles ready: Investor wallet, Rentor wallet, Renter wallet",
    "MongoDB running with seed data loaded",
    "Screen recorder ready (OBS / QuickTime / Loom)",
    "Resolution set to 1920x1080 for clear recording",
]
for item in checklist:
    p = doc.add_paragraph(item, style="List Bullet")

add_blank()

# ── Timing Overview ─────────────────────────────────────────────────────────

add_section("Timing Overview (< 5 Minutes)")

table = doc.add_table(rows=8, cols=3)
table.style = "Light Shading Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Segment", "Duration", "Cumulative"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

rows_data = [
    ("Opening — Landing Page & Value Prop", "0:30", "0:30"),
    ("Investor — Onboarding + WorldID + Invest", "1:15", "1:45"),
    ("Rentor — Vehicle + Fundraising + CRE", "1:00", "2:45"),
    ("Renter — Browse + Book", "0:30", "3:15"),
    ("Admin — KYC + CRE Activity + Milestones", "0:45", "4:00"),
    ("CRE Workflow Demo (Live UI)", "0:45", "4:45"),
    ("Closing — Summary", "0:15", "5:00"),
]
for i, (seg, dur, cum) in enumerate(rows_data):
    table.rows[i + 1].cells[0].text = seg
    table.rows[i + 1].cells[1].text = dur
    table.rows[i + 1].cells[2].text = cum

add_blank()
add_divider()

# ═══════════════════════════════════════════════════════════════════════════
# SEGMENT 1: OPENING
# ═══════════════════════════════════════════════════════════════════════════

add_section("SEGMENT 1: Opening — Landing Page  [0:00 – 0:30]")

add_step(1, "Landing Page Hero", "0:00")
add_action("Open browser at http://localhost:3000 — landing page loads")
add_say("RegShield is a tokenized vehicle rental platform. Investors fund real vehicles through milestone-based escrow, receive ERC-3643 security tokens, and earn quarterly rental revenue — all verified on-chain.")
add_show("Hero section with orbiting badges: AssetToken, RevenueToken, Escrow, CRE Verified")
add_blank()

add_step(2, "How It Works + Sponsor Highlights", "0:15")
add_action("Scroll down slowly to 'How It Works' section")
add_say("The platform uses two key sponsor technologies. Chainlink CRE automates milestone verification — VIN validation, purchase confirmation, insurance, and registration checks — replacing centralized backend jobs. WorldID provides proof-of-personhood to prevent sybil attacks on our investor onboarding.")
add_show("4-step flow: Invest → CRE Verifies → Tokens Minted → Earn Revenue")
add_tip("Pause briefly on the 'Chainlink CRE Automation' and 'OnchainID Identity' feature cards")
add_blank()

add_divider()

# ═══════════════════════════════════════════════════════════════════════════
# SEGMENT 2: INVESTOR
# ═══════════════════════════════════════════════════════════════════════════

add_section("SEGMENT 2: Investor Flow  [0:30 – 1:45]")

add_step(3, "Investor Onboarding — Connect Wallet", "0:30")
add_action("Click 'Start Investing' on hero OR navigate to /onboarding")
add_action("Select 'Investor' role card")
add_say("Let's walk through the investor experience. A new user selects the Investor role and begins the verification process.")
add_show("Role selection: Investor (blue), Rentor (green), Renter (purple)")
add_blank()

add_step(4, "Wallet Connect + Bind", "0:40")
add_action("Click 'Connect Wallet' — MetaMask popup appears → approve")
add_action("Click 'Bind Wallet' to link wallet to account")
add_say("The investor connects their MetaMask wallet and binds it to their account. This wallet address becomes their on-chain identity.")
add_show("Wallet connected indicator with truncated address")
add_blank()

add_step(5, "WorldID Verification (Sponsor Highlight)", "0:50")
add_action("Show the WorldID verification step on the verification page")
add_say("Here's where WorldID comes in. We integrate WorldID's proof-of-personhood to prevent sybil attacks. Each investor must be a unique human — verified through WorldID's Orb-level zero-knowledge proof. This is verified on-chain through our WorldIDVerifier contract on Sepolia.")
add_show("WorldIDVerifyButton component — either 'Verified' badge or 'Open in World App' info panel")
add_note("If not in World App, the UI shows an info banner explaining to open RegShield inside World App. For demo, you can show the verified state if the wallet is already verified, or show the UI explaining the flow.")
add_tip("Emphasize: 'WorldID verification is stored on-chain — our smart contract WorldIDVerifier.sol calls WorldID's router contract to verify the ZK proof, then maps the address as verified. This prevents one person from creating multiple investor accounts.'")
add_blank()

add_step(6, "Investor Type Selection", "1:05")
add_action("Select 'Retail Investor' (recommended) on the investor type step")
add_say("Investors choose their tier — Retail, Accredited, or Institutional — each with different investment limits and lock periods. The type is registered on-chain via our InvestorTypeRegistry contract for ERC-3643 compliance.")
add_show("Three tier cards with lock amounts, limits, and lock periods")
add_blank()

add_step(7, "KYC Upload", "1:10")
add_action("Show the KYC document upload step (primary ID + proof of address)")
add_say("KYC documents are submitted and reviewed. Once approved by admin, the investor's identity is registered on-chain in our ERC-3643 IdentityRegistry — enabling compliant token transfers.")
add_show("Upload fields for Primary Document and Proof of Address")
add_blank()

add_step(8, "Marketplace — Invest in a Vehicle", "1:20")
add_action("Navigate to /investor/marketplace")
add_action("Click on a vehicle campaign card to open the investment modal")
add_say("On the marketplace, the investor browses active fundraising campaigns. Each campaign shows the vehicle, target amount, current progress, and expected ROI. Prices are shown in both ETH and USD using Chainlink's ETH/USD price feed on Sepolia.")
add_show("Campaign cards with progress bars + ETH/USD prices via Chainlink")
add_tip("Highlight the ETH/USD conversion: 'These real-time USD conversions come from Chainlink's price feed oracle on Sepolia.'")
add_blank()

add_step(9, "Investment Confirmation", "1:35")
add_action("Enter investment amount in the modal → click Invest → MetaMask tx popup")
add_say("The investor commits ETH which goes into our PaymentEscrow contract. Funds stay locked until Chainlink CRE verifies all four milestones — we'll see that in the CRE demo later.")
add_show("InvestmentModal with amount input + MetaMask confirmation")
add_blank()

add_divider()

# ═══════════════════════════════════════════════════════════════════════════
# SEGMENT 3: RENTOR
# ═══════════════════════════════════════════════════════════════════════════

add_section("SEGMENT 3: Rentor Flow  [1:45 – 2:45]")

add_step(10, "Switch to Rentor Account", "1:45")
add_action("Log out or use Role Switcher → switch to Rentor account")
add_action("Navigate to /rentor/dashboard")
add_say("Now let's see the vehicle owner's perspective. The Rentor lists vehicles, creates fundraising campaigns, and manages bookings.")
add_show("Rentor dashboard with vehicle count, bookings, and monthly revenue")
add_blank()

add_step(11, "My Vehicles — NFT-Registered Fleet", "1:55")
add_action("Navigate to /rentor/vehicles")
add_action("Click on a vehicle card to show details")
add_say("Each vehicle is registered as an ERC-721 NFT on our VehicleNFT contract. This stores the VIN, make, model, mileage, insurance expiry, and registration expiry — all on-chain. The Chainlink CRE workflows read this data to verify milestones automatically.")
add_show("Vehicle cards with status badges (Available / Fundraising)")
add_blank()

add_step(12, "Fundraising Campaign", "2:05")
add_action("Navigate to /rentor/fundraising")
add_action("Show an active campaign with investors")
add_say("The Rentor creates a fundraising campaign for each vehicle. Once the target is reached and Chainlink CRE verifies all four milestones — VIN validation, purchase confirmation, insurance check, and registration check — escrow funds are released and ERC-3643 tokens are minted to investors.")
add_show("Campaign dashboard with target vs raised, investor count")
add_tip("Emphasize: 'This is where Chainlink CRE replaces manual admin verification — the 4-milestone check happens automatically through our CRE workflows.'")
add_blank()

add_step(13, "Revenue & Analytics", "2:20")
add_action("Navigate to /rentor/analytics")
add_say("The Rentor tracks revenue from vehicle rentals. Revenue follows a waterfall distribution — 50% to token-holding investors, 10% operator fee to the rentor, remainder to reserves. The ETH/USD conversion here uses Chainlink's price feed.")
add_show("Revenue chart with period selector + vehicle stats")
add_blank()

add_step(14, "Bookings Management", "2:35")
add_action("Navigate to /rentor/bookings → show a pending booking")
add_say("Incoming rental bookings can be approved or denied. In the CRE version, our onboarding workflow auto-checks renter compliance before approval.")
add_show("Booking cards with approve/deny actions")
add_blank()

add_divider()

# ═══════════════════════════════════════════════════════════════════════════
# SEGMENT 4: RENTER
# ═══════════════════════════════════════════════════════════════════════════

add_section("SEGMENT 4: Renter Flow  [2:45 – 3:15]")

add_step(15, "Switch to Renter Account", "2:45")
add_action("Switch to Renter account → navigate to /renter/browse")
add_say("The Renter browses available vehicles, books them, and leaves reviews. Let's do a quick walkthrough.")
add_show("Vehicle browse page with filters (category, price, location)")
add_blank()

add_step(16, "Browse & Book a Vehicle", "2:55")
add_action("Apply a filter (e.g., price range or category)")
add_action("Click a vehicle → show detail page → click Book")
add_say("The renter can filter by price, category, fuel type, and location. After selecting a vehicle, they see full details, reviews, and availability. Booking is straightforward — select dates and confirm.")
add_show("Vehicle detail page with images, specs, reviews, and booking calendar")
add_blank()

add_step(17, "My Bookings", "3:05")
add_action("Navigate to /renter/bookings")
add_say("The renter tracks all their bookings and can leave reviews after completing a rental.")
add_show("Booking list with status filters (Active, Past)")
add_blank()

add_divider()

# ═══════════════════════════════════════════════════════════════════════════
# SEGMENT 5: ADMIN
# ═══════════════════════════════════════════════════════════════════════════

add_section("SEGMENT 5: Admin Panel  [3:15 – 4:00]")

add_step(18, "Admin Dashboard Overview", "3:15")
add_action("Switch to Admin account → navigate to /admin")
add_say("The admin panel provides full oversight of the platform — KYC reviews, milestone tracking, revenue management, and most importantly, the CRE activity feed.")
add_show("Admin dashboard with pending KYC count, vehicle registrations, platform stats")
add_blank()

add_step(19, "KYC Management", "3:25")
add_action("Navigate to /admin/kyc → click a pending submission → approve it")
add_say("Admins review KYC submissions. On approval, the investor's identity is registered on-chain in our ERC-3643 IdentityRegistry contract, enabling compliant token transfers.")
add_show("KYC submission detail with document preview and approve/reject buttons")
add_blank()

add_step(20, "Milestone Tracking (Chainlink CRE)", "3:35")
add_action("Navigate to /admin/milestones")
add_say("This page shows the 4-milestone verification status for each vehicle investment. In production, Chainlink CRE workflows automatically verify each milestone — VIN validation via the NHTSA API, purchase escrow confirmation, insurance expiry check, and registration check — all from on-chain VehicleNFT data.")
add_show("Milestone grid: Vehicle ID → 4 checkmarks (VIN, Purchase, Insurance, Registration)")
add_tip("Emphasize: 'Each of these milestones maps to a CRE workflow — we'll see those workflows in action next on the CRE Activity Feed.'")
add_blank()

add_step(21, "CRE Activity Feed", "3:50")
add_action("Navigate to /admin/cre-activity")
add_say("The CRE activity feed shows real-time events from Chainlink's Compute Runtime Environment — every verification, every milestone check, every compliance action logged with timestamps. Let me show you what this looks like when the CRE workflows are actively running.")
add_show("Activity feed page with workflow filter tabs and 'Simulate Events' button in header")
add_blank()

add_divider()

# ═══════════════════════════════════════════════════════════════════════════
# SEGMENT 6: CRE LIVE UI DEMO
# ═══════════════════════════════════════════════════════════════════════════

add_section("SEGMENT 6: CRE Workflow Demo  [4:00 – 4:45]")

add_step(22, "Activate CRE Simulation", "4:00")
add_action("On the CRE Activity Feed page (/admin/cre-activity), click the 'Simulate Events' button in the Events section header")
add_say("Now let's see the Chainlink CRE workflows in action. We have 5 workflows running on the Chainlink DON that replace centralized backend cron jobs with decentralized, verifiable automation. I'll activate the simulation to show what these events look like as they flow through the system.")
add_show("Amber 'Demo Mode Active' banner appears at the top. The activity feed populates with 10 simulated CRE events across all 5 workflow types, each tagged with a 'Simulated' label.")
add_blank()

add_step(23, "Walk Through Compliance & Vehicle Events", "4:10")
add_action("Scroll through the simulated events — point out the Compliance and Vehicle workflow events")
add_say("The compliance workflow detects expired registration and insurance on vehicles by reading on-chain VehicleNFT data. Here you can see it flagged an expired registration and triggered a suspension. The vehicle workflow monitors telematics — updating mileage and recording maintenance events directly on-chain.")
add_show("Events showing: 'Registration Renewed' (compliance), 'Vehicle Suspended' (compliance), 'Mileage Updated' (vehicle), 'Maintenance Recorded' (vehicle)")
add_tip("Point to the event details: contract address, block number, and timestamp — emphasize these are real on-chain events when running in production")
add_blank()

add_step(24, "Walk Through Payment & Onboarding Events", "4:25")
add_action("Click the 'Payment' and 'Onboarding' filter tabs to isolate those event types")
add_say("The payment workflow handles revenue distribution — splitting rental income between investors, operators, and reserves according to the waterfall rules. The onboarding workflow auto-verifies new investors by checking their WorldID status and KYC completion, then submitting milestone reports to release escrowed funds.")
add_show("Payment events: 'Revenue Distributed', 'Payment Processed'. Onboarding events: 'Investor Verified', 'KYC Processed'. Each event shows the receiver contract address and block details.")
add_blank()

add_step(25, "Highlight: 5 Workflows Replace Backend", "4:35")
add_action("Click the 'All Workflows' tab to show all events together")
add_say("In total, RegShield has 5 CRE workflows that replace centralized backend services: Onboarding replaces manual admin approval. Campaign Monitor replaces campaignScheduler.js. Rental Lifecycle replaces bookingScheduler.js. Vehicle Telematics replaces revenueSyncService.js. And Compliance Monitoring is an entirely new capability that only CRE makes possible — detecting expired registration, insurance, and overdue maintenance automatically.")
add_show("Full activity feed with all 10 events from 5 workflow types — compliance (green), vehicle (blue), payment (purple), onboarding (cyan), campaign (amber)")
add_tip("Keep this concise — list them quickly without pausing on each")
add_blank()

add_divider()

# ═══════════════════════════════════════════════════════════════════════════
# SEGMENT 7: CLOSING
# ═══════════════════════════════════════════════════════════════════════════

add_section("SEGMENT 7: Closing  [4:45 – 5:00]")

add_step(26, "Summary & Sponsors", "4:45")
add_action("Switch back to browser — show landing page")
add_say("To summarize: RegShield uses Chainlink CRE to replace 4 centralized backend services and add 1 new capability — all with decentralized, verifiable automation. WorldID provides sybil-resistant identity verification for investors. Together with ERC-3643 security tokens and milestone-based escrow, we've built a fully on-chain compliant vehicle investment platform. Thank you.")
add_show("Landing page with sponsor technology highlights visible")
add_blank()

add_divider()

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════════════════════

add_section("Appendix A: Sponsor Technology Summary")

table2 = doc.add_table(rows=4, cols=3)
table2.style = "Light Shading Accent 1"

headers2 = ["Sponsor", "Integration", "Where Shown"]
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

sponsor_data = [
    ("Chainlink CRE", "5 off-chain workflows: onboarding, campaign, rental, vehicle, compliance. Replace centralized cron jobs with decentralized automation.", "Admin milestones page, CRE activity feed (live events + simulation mode)"),
    ("Chainlink Price Feed", "ETH/USD oracle (Sepolia) for real-time price conversion throughout platform.", "Marketplace, dashboards, investment modal, revenue analytics"),
    ("WorldID", "Proof-of-personhood via Orb-level ZK verification. On-chain via WorldIDVerifier.sol. Prevents sybil attacks on investor onboarding.", "Verification page, onboarding flow, WorldIDVerifyButton component"),
]
for i, (sponsor, integration, shown) in enumerate(sponsor_data):
    table2.rows[i + 1].cells[0].text = sponsor
    table2.rows[i + 1].cells[1].text = integration
    table2.rows[i + 1].cells[2].text = shown

add_blank()

add_section("Appendix B: CRE Workflow → Backend Replacement Map")

table3 = doc.add_table(rows=6, cols=3)
table3.style = "Light Shading Accent 1"

headers3 = ["CRE Workflow", "Replaces", "Reports Submitted"]
for i, h in enumerate(headers3):
    table3.rows[0].cells[i].text = h
    for p in table3.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

wf_data = [
    ("onboarding-workflow", "Manual admin approval", "5 onboarding reports"),
    ("campaign-workflow", "campaignScheduler.js", "2 campaign reports"),
    ("rental-workflow", "bookingScheduler.js + milestones", "3 milestone reports"),
    ("vehicle-workflow", "revenueSyncService.js (partial)", "3 telemetry reports"),
    ("compliance-workflow", "None (new CRE capability)", "3 compliance reports"),
]
for i, (wf, replaces, reports) in enumerate(wf_data):
    table3.rows[i + 1].cells[0].text = wf
    table3.rows[i + 1].cells[1].text = replaces
    table3.rows[i + 1].cells[2].text = reports

add_blank()

add_section("Appendix C: Key Smart Contracts (Sepolia)")

table4 = doc.add_table(rows=9, cols=2)
table4.style = "Light Shading Accent 1"

headers4 = ["Contract", "Address"]
for i, h in enumerate(headers4):
    table4.rows[0].cells[i].text = h
    for p in table4.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

contracts = [
    ("WorldIDVerifier", "0x838C9397F5c00f7010924dDc4c2E93Fcab6c0363"),
    ("IdentityRegistry", "0x188C09768E0b2D21212FCDb0faEef175e55d927A"),
    ("VehicleNFT", "0xfcE0FD3671E99D65E0FF70b30b9238bB83D91814"),
    ("OnboardingReceiver", "0xF080a8B7Ee2e83c9beE26a795e43D70b1D093850"),
    ("CampaignMonitorReceiver", "0x84A9B21B7d2Ba6120923edAA32B283fD2E35FB94"),
    ("ComplianceReceiver", "0x0eA9cd084287107BCA0f9785B030c22Db72301fD"),
    ("VehicleReceiver", "0x73c58B5Ba299FaAA64103E453ba55b408C91e81B"),
    ("PaymentReceiver", "0x2f7F8ED26B72A43988AfA1f3088Bd4969f39B7C2"),
]
for i, (name, addr) in enumerate(contracts):
    table4.rows[i + 1].cells[0].text = name
    table4.rows[i + 1].cells[1].text = addr

add_blank()

add_section("Appendix D: Quick Troubleshooting")

troubles = [
    ("MetaMask not connecting", "Ensure Sepolia testnet is selected. Clear site data if wallet state is stale."),
    ("WorldID shows 'Not in World App'", "Expected on desktop browser. For demo, show the UI state and explain the World App flow."),
    ("CRE simulated events not showing", "Click 'Simulate Events' button on /admin/cre-activity. If button is missing, ensure frontend is running the latest build."),
    ("Frontend shows $0 revenue", "Run 'node backend/services/revenueSyncService.js' to trigger manual sync."),
    ("KYC page empty", "Ensure MongoDB is running with seed data. Check backend console for errors."),
]
for problem, solution in troubles:
    p = doc.add_paragraph()
    run = p.add_run(f"{problem}: ")
    run.bold = True
    run = p.add_run(solution)

# ── Save ────────────────────────────────────────────────────────────────────

output_path = "/Volumes/extreme/Projects/Solidity/Portfolio/RegShield/rental-cre/demo/RegShield_Demo_Script.docx"
doc.save(output_path)
print(f"Demo script saved to: {output_path}")
